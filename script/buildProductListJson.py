import docx
import json
import os
import re
import traceback

file = r'./production description.docx'
imagePath = r'../website/public/assets/images/products/'
iConPath = r'../website/public/assets/images/product-icon/'
doc = docx.Document(file)

productList = {}
errorList = {"missImage": [], "missIcon": []}
stk = [productList,]


def checkPageExists(path, pid):
    for suffix in ['jpg', 'png', 'jpeg']:
        filePath = os.path.join(path, f'{pid}.{suffix}')
        if os.path.exists(filePath):
            return filePath
    return None


def startWithOrder(text):
    return re.search('^\d+\.', text) is not None


for idx, paragraph in enumerate(doc.paragraphs):
    try:
        if paragraph.style.name.startswith('Heading'):
            num = int(paragraph.style.name.replace('Heading', ''))
            assert (num <= len(stk))
            while num < len(stk):
                stk.pop()
            text = paragraph.text.strip().strip(':')
            stk[-1][text] = {} if num < 3 else []
            stk.append(stk[-1][text])
        else:
            if len(paragraph.text.strip()) == 0:
                continue
            if len(stk[-1]) > 0 and startWithOrder(stk[-1][-1]) and not startWithOrder(paragraph.text.strip()):
                raise Exception('Invalid paragraph')
            stk[-1].append(paragraph.text.strip())
    except Exception:
        print(traceback.format_exc())
        pLen = len(doc.paragraphs)
        for i in range(idx, idx + 10):
            if i < pLen:
                print(doc.paragraphs[i].text)
        break

procceededProductList = {}

for category_name, products in productList.items():
    for product_name, info in products.items():
        info['pdfUrl'] = '#'
        info['desc'] = []
        for key in list(info):
            if 'Other Names'.lower() in key.lower():
                info.pop(key)
            if 'Description'.lower() in key.lower():
                tmp = info.pop(key)
                info['desc'] = tmp
        pid = '_'.join(product_name.lower().split(' '))
        category_idx = '_'.join(category_name.lower().split(' '))
        if pid not in procceededProductList:
            procceededProductList[pid] = {"pid": pid, "title": product_name, "categories": [], "category_idx": [], "infos": {}}
        procceededProductList[pid]["categories"].append(category_name)
        procceededProductList[pid]['category_idx'].append(category_idx)
        procceededProductList[pid]["infos"][category_idx] = info
        if checkPageExists(imagePath, pid) is not None:
            procceededProductList[pid]['image'] = checkPageExists(imagePath, pid)
        else:
            errorList["missImage"].append(pid)
            procceededProductList[pid]['image'] = "/assets/images/gallery/gallery-1.jpg"

        if checkPageExists(iConPath, pid) is not None:
            procceededProductList[pid]['overlay'] = checkPageExists(iConPath, pid)
        else:
            errorList["missIcon"].append(pid)
            procceededProductList[pid]['overlay'] = "/assets/images/icons/icon-8.png"

with open('error_list.json', 'w') as f:
    json.dump(errorList, f, ensure_ascii=False, indent=4)

with open('product.json', 'w', encoding='utf8') as f:
    json.dump(list(procceededProductList.values()), f, ensure_ascii=False, indent=4)